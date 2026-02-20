"""
Router pour la comparaison entre l'analyse humaine et l'analyse IA.
Permet de confronter l'évaluation de l'utilisateur avec celle de l'IA.
"""

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from io import BytesIO
from docx import Document
from pydantic import BaseModel, Field
from typing import Optional, List
from app.services.llm_service import call_llm_for_risk
from app.services.ml_service import predict_classification
from app.utils.normalize import kinney_score, classify_from_score
from app.constants import RISK_CATEGORIES, RISK_TYPES

router = APIRouter(tags=["Comparaison Humain vs IA"])


class CompareRequest(BaseModel):
    """Requête de comparaison humain vs IA."""
    description: str = Field(..., description="Description du risque à analyser")
    category: str = Field(..., description="Catégorie du risque")
    type: str = Field(..., description="Type de risque")
    sector: Optional[str] = Field("", description="Secteur d'activité")

    # Évaluation humaine
    user_G: int = Field(..., ge=1, le=5, description="Gravité évaluée par l'humain (1-5)")
    user_F: int = Field(..., ge=1, le=5, description="Fréquence évaluée par l'humain (1-5)")
    user_P: int = Field(..., ge=1, le=5, description="Probabilité évaluée par l'humain (1-5)")
    user_classification: Optional[str] = Field(
        None,
        description="Classification manuelle par l'humain (optionnel)"
    )

    model_config = {
        "json_schema_extra": {
            "example": {
                "description": "Risque de cyberattaque sur le système de paiement en ligne",
                "category": "Industriel",
                "type": "Cyber & SSI",
                "sector": "Technologie",
                "user_G": 4,
                "user_F": 3,
                "user_P": 3,
                "user_classification": "Moyen"
            }
        }
    }


class AnalysisDetail(BaseModel):
    """Détail d'une analyse (humaine ou IA)."""
    G: int
    F: int
    P: int
    score: int
    classification: str


class CompareResponse(BaseModel):
    """Réponse de la comparaison humain vs IA."""
    description: str = Field(..., description="Description du risque analysé")
    category: str = Field(..., description="Catégorie")
    type: str = Field(..., description="Type")

    # Analyse humaine
    human_analysis: AnalysisDetail = Field(..., description="Analyse réalisée par l'humain")

    # Analyse IA avec causes et recommandations détaillées
    ia_analysis: IAAnalysisDetail = Field(..., description="Analyse réalisée par l'IA avec causes et recommandations")

    # Comparaison
    comparison: dict = Field(..., description="Résultat de la comparaison")


class ProjectAnalysisRequest(BaseModel):
    """Requête d'analyse complète d'un projet par l'IA."""
    description: str = Field(..., description="Description complète du projet")
    sector: Optional[str] = Field("", description="Secteur d'activité")
    entity_services: Optional[str] = Field("", description="Services de l'entité")
    analysis_title: str = Field(..., description="Titre de l'analyse")

class IdentifiedRisk(BaseModel):
    """Risque identifié par l'IA."""
    title: str = Field(..., description="Titre du risque")
    description: str = Field(..., description="Description détaillée du risque")
    category: str = Field(..., description="Catégorie (Projet/Programme, Industriel, Qualité)")
    type: str = Field(..., description="Type (Commercial, Financier, Technique, Cyber & SSI)")
    estimated_G: int = Field(..., ge=1, le=5, description="Gravité estimée (1-5)")
    estimated_F: int = Field(..., ge=1, le=5, description="Fréquence estimée (1-5)")
    estimated_P: int = Field(..., ge=1, le=5, description="Probabilité estimée (1-5)")

class ProjectAnalysisResponse(BaseModel):
    """Réponse de l'analyse IA du projet."""
    project_title: str
    identified_risks: List[IdentifiedRisk] = Field(..., description="Liste des risques identifiés par l'IA")
    general_recommendations: List[str] = Field(default=[], description="Recommandations générales pour le projet")


def calculate_difference(human_val: int, ia_val: int) -> dict:
    """Calcule la différence entre valeurs humaine et IA."""
    diff = human_val - ia_val
    if diff == 0:
        assessment = "identique"
    elif diff > 0:
        assessment = "humain plus élevé"
    else:
        assessment = "IA plus élevée"

    return {
        "human": human_val,
        "ia": ia_val,
        "difference": abs(diff),
        "assessment": assessment
    }


@router.post("/compare", response_model=CompareResponse)
async def compare_human_vs_ia(request: CompareRequest):
    """
    Compare l'analyse de risque humaine avec l'analyse IA.

    Cette endpoint permet de :
    1. Soumettre l'évaluation humaine (G, F, P)
    2. Obtenir l'évaluation de l'IA pour le même risque
    3. Comparer les deux analyses point par point
    4. Identifier les divergences et leurs causes possibles

    Utile pour :
    - Calibrer les évaluations humaines
    - Identifier les biais d'évaluation
    - Former les analystes de risques
    - Valider les analyses critiques
    """
    # Validation des entrées
    if request.category not in RISK_CATEGORIES:
        raise HTTPException(
            status_code=400,
            detail=f"Catégorie invalide. Valeurs acceptées: {RISK_CATEGORIES}"
        )

    if request.type not in RISK_TYPES:
        raise HTTPException(
            status_code=400,
            detail=f"Type invalide. Valeurs acceptées: {RISK_TYPES}"
        )

    # Calcul analyse humaine
    human_score = kinney_score(request.user_G, request.user_F, request.user_P)
    human_classification = request.user_classification or classify_from_score(human_score)

    # Analyse IA
    llm_out = call_llm_for_risk(request.description, request.category, request.type, request.sector)

    if llm_out is None:
        raise HTTPException(status_code=503, detail="Service IA indisponible. Réessayez plus tard.")

    try:
        ia_G = int(llm_out.get("G"))
        ia_F = int(llm_out.get("F"))
        ia_P = int(llm_out.get("P"))
        ia_causes = llm_out.get("causes", [])
        ia_recommendations = llm_out.get("recommendations", [])
        ia_justification = llm_out.get("justification", "")
    except Exception:
        raise HTTPException(status_code=502, detail="Réponse IA malformée")

    ia_score = kinney_score(ia_G, ia_F, ia_P)
    ia_classification = classify_from_score(ia_score)

    # Comparaison détaillée
    g_comparison = calculate_difference(request.user_G, ia_G)
    f_comparison = calculate_difference(request.user_F, ia_F)
    p_comparison = calculate_difference(request.user_P, ia_P)
    score_comparison = calculate_difference(human_score, ia_score)

    # Déterminer l'accord global
    classifications_match = human_classification == ia_classification
    scores_close = abs(human_score - ia_score) <= 10  # Tolérance de 10 points

    if classifications_match and scores_close:
        agreement_level = "fort"
        agreement_message = "L'analyse humaine et l'IA sont en accord."
    elif classifications_match:
        agreement_level = "modéré"
        agreement_message = "Les classifications concordent mais les scores diffèrent significativement."
    elif scores_close:
        agreement_level = "modéré"
        agreement_message = "Les scores sont proches mais les classifications diffèrent (zone limite)."
    else:
        agreement_level = "faible"
        agreement_message = "Divergence significative entre l'analyse humaine et l'IA. Une revue est recommandée."

    # Identifier le facteur de plus grande divergence
    max_diff_factor = max(
        [("G (Gravité)", g_comparison["difference"]),
         ("F (Fréquence)", f_comparison["difference"]),
         ("P (Probabilité)", p_comparison["difference"])],
        key=lambda x: x[1]
    )

    comparison = {
        "G": g_comparison,
        "F": f_comparison,
        "P": p_comparison,
        "score": score_comparison,
        "classifications_match": classifications_match,
        "agreement_level": agreement_level,
        "agreement_message": agreement_message,
        "max_divergence_factor": max_diff_factor[0] if max_diff_factor[1] > 0 else None,
        "recommendations": []
    }

    # Ajouter des recommandations basées sur les divergences
    if not classifications_match:
        if human_classification == "Faible" and ia_classification in ["Moyen", "Eleve"]:
            comparison["recommendations"].append(
                "⚠️ L'IA évalue ce risque plus sévèrement. Vérifiez si certains impacts n'ont pas été sous-estimés."
            )
        elif human_classification == "Eleve" and ia_classification in ["Faible", "Moyen"]:
            comparison["recommendations"].append(
                "💡 L'IA évalue ce risque moins sévèrement. Vérifiez si des mesures de mitigation existantes n'ont pas été prises en compte."
            )

    if g_comparison["difference"] >= 2:
        comparison["recommendations"].append(
            f"📊 Écart important sur la Gravité ({g_comparison['difference']} points). Revoyez l'impact potentiel."
        )

    if f_comparison["difference"] >= 2:
        comparison["recommendations"].append(
            f"📊 Écart important sur la Fréquence ({f_comparison['difference']} points). Revoyez la fréquence d'exposition."
        )

    if p_comparison["difference"] >= 2:
        comparison["recommendations"].append(
            f"📊 Écart important sur la Probabilité ({p_comparison['difference']} points). Revoyez la vraisemblance."
        )

    return CompareResponse(
        description=request.description,
        category=request.category,
        type=request.type,
        human_analysis=AnalysisDetail(
            G=request.user_G,
            F=request.user_F,
            P=request.user_P,
            score=human_score,
            classification=human_classification
        ),
        ia_analysis=IAAnalysisDetail(
            G=ia_G,
            F=ia_F,
            P=ia_P,
            score=ia_score,
            classification=ia_classification,
            causes=ia_causes,
            recommendations=ia_recommendations,
            justification=ia_justification
        ),
        comparison=comparison
    )


@router.post("/compare/report")
async def export_compare_report(request: CompareRequest):
    """
    Génère un document Word (.docx) contenant la comparaison Humain vs IA
    avec tableaux de valeurs et représentations simples des graphes (barres textuelles).
    """
    # Reutiliser la logique de comparaison
    # Validation
    if request.category not in RISK_CATEGORIES:
        raise HTTPException(status_code=400, detail=f"Catégorie invalide. Valeurs acceptées: {RISK_CATEGORIES}")
    if request.type not in RISK_TYPES:
        raise HTTPException(status_code=400, detail=f"Type invalide. Valeurs acceptées: {RISK_TYPES}")

    human_score = kinney_score(request.user_G, request.user_F, request.user_P)
    human_classification = request.user_classification or classify_from_score(human_score)

    llm_out = call_llm_for_risk(request.description, request.category, request.type, request.sector)
    if llm_out is None:
        raise HTTPException(status_code=503, detail="Service IA indisponible. Réessayez plus tard.")
    try:
        ia_G = int(llm_out.get("G"))
        ia_F = int(llm_out.get("F"))
        ia_P = int(llm_out.get("P"))
        ia_causes = llm_out.get("causes", [])
        ia_recommendations = llm_out.get("recommendations", [])
        ia_justification = llm_out.get("justification", "")
    except Exception:
        raise HTTPException(status_code=502, detail="Réponse IA malformée")

    ia_score = kinney_score(ia_G, ia_F, ia_P)
    ia_classification = classify_from_score(ia_score)
    
    # Normaliser les scores sur 100
    human_score_normalized = int(round(human_score / 125 * 100))
    ia_score_normalized = int(round(ia_score / 125 * 100))

    # Document
    doc = Document()
    doc.add_heading("Comparaison Humain vs IA", level=0)
    doc.add_paragraph(f"Description: {request.description}")
    doc.add_paragraph(f"Catégorie: {request.category}  |  Type: {request.type}  |  Secteur: {request.sector or ''}")

    doc.add_heading("Analyse Utilisateur", level=1)
    doc.add_paragraph(f"G: {request.user_G}  F: {request.user_F}  P: {request.user_P}")
    doc.add_paragraph(f"Score: {human_score_normalized} / 100  |  Classification: {human_classification}")

    doc.add_heading("Analyse IA", level=1)
    doc.add_paragraph(f"G: {ia_G}  F: {ia_F}  P: {ia_P}")
    doc.add_paragraph(f"Score: {ia_score_normalized} / 100  |  Classification: {ia_classification}")
    if ia_justification:
        doc.add_paragraph(f"Justification IA: {ia_justification}")
    if ia_causes:
        doc.add_paragraph("Causes: " + "; ".join(map(str, ia_causes)))
    if ia_recommendations:
        doc.add_paragraph("Recommandations: " + "; ".join(map(str, ia_recommendations)))

    # Graphiques simplifiés (barres textuelles)
    doc.add_heading("Graphiques comparatifs (G/F/P)", level=1)
    def bar(val: int, max_val: int = 5, length: int = 20) -> str:
        pct = max(0, min(max_val, val)) / float(max_val)
        filled = int(round(pct * length))
        return "█" * filled + "░" * (length - filled)

    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    table.cell(0,0).text = "Facteur"
    table.cell(0,1).text = "Vous"
    table.cell(0,2).text = "IA"

    # G
    table.cell(1,0).text = "G"
    table.cell(1,1).text = f"{request.user_G}  |  {bar(request.user_G)}"
    table.cell(1,2).text = f"{ia_G}  |  {bar(ia_G)}"
    # F
    table.cell(2,0).text = "F"
    table.cell(2,1).text = f"{request.user_F}  |  {bar(request.user_F)}"
    table.cell(2,2).text = f"{ia_F}  |  {bar(ia_F)}"
    # P
    table.cell(3,0).text = "P"
    table.cell(3,1).text = f"{request.user_P}  |  {bar(request.user_P)}"
    table.cell(3,2).text = f"{ia_P}  |  {bar(ia_P)}"

    # Graphique du score global
    doc.add_heading("Score Global (sur 100)", level=1)
    
    def score_bar(val: int, max_val: int = 100, length: int = 50) -> str:
        pct = max(0, min(max_val, val)) / float(max_val)
        filled = int(round(pct * length))
        return "█" * filled + "░" * (length - filled)
    
    score_table = doc.add_table(rows=3, cols=2)
    score_table.style = "Table Grid"
    score_table.cell(0,0).text = "Analyse"
    score_table.cell(0,1).text = "Score / 100"
    
    score_table.cell(1,0).text = "Utilisateur"
    score_table.cell(1,1).text = f"{human_score_normalized}/100  |  {score_bar(human_score_normalized)}"
    
    score_table.cell(2,0).text = "IA"
    score_table.cell(2,1).text = f"{ia_score_normalized}/100  |  {score_bar(ia_score_normalized)}"
    
    # Écarts
    doc.add_heading("Comparaison & Écarts", level=1)
    doc.add_paragraph(f"Écart score: {abs(human_score_normalized - ia_score_normalized)} points (sur 100)")
    doc.add_paragraph("Les graphiques ci-dessus sont indicatifs (barres textuelles).")

    # Stream
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    headers = {"Content-Disposition": "attachment; filename=compare_report.docx"}
    return StreamingResponse(bio, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers=headers)


@router.post("/analyze/project", response_model=ProjectAnalysisResponse)
async def analyze_project_with_ai(request: ProjectAnalysisRequest):
    """
    Analyse complète d'un projet par l'IA pour identifier dynamiquement les risques.

    Au lieu d'utiliser des risques statiques basés sur des mots-clés, cette endpoint:
    1. Analyse la description complète du projet
    2. Identifie les risques spécifiques au contexte
    3. Fournit des estimations G/F/P pour chaque risque
    4. Génère des recommandations générales

    Cela permet une analyse vraiment personnalisée et pertinente.
    """
    # Prompt pour l'IA
    prompt = f"""
Tu es un expert en gestion des risques utilisant la méthode Kinney. Analyse ce projet et identifie les risques spécifiques.

## Contexte du projet
- Titre: {request.analysis_title}
- Description: {request.description}
- Secteur: {request.sector or 'Non spécifié'}
- Services: {request.entity_services or 'Non spécifiés'}

## Échelles Kinney
- Gravité (G): 1=Négligeable, 2=Faible, 3=Modérée, 4=Grave, 5=Catastrophique
- Fréquence (F): 1=Rare, 2=Occasionnel, 3=Fréquent, 4=Très fréquent, 5=Permanent
- Probabilité (P): 1=Improbable, 2=Peu probable, 3=Probable, 4=Très probable, 5=Quasi certain

## Catégories disponibles
- Projet/Programme: risques liés à la gestion et exécution du projet
- Industriel: risques techniques et opérationnels
- Qualité: risques liés à la conformité, sécurité des données, qualité

## Types disponibles
- Commercial: risques business et stratégiques
- Financier: risques budgétaires et économiques
- Technique: risques technologiques et opérationnels
- Cyber & SSI: risques de sécurité informatique et données

## Instructions
Identifie 4-6 risques SPECIFIQUES à ce projet. Pour chaque risque:
1. Un titre concis et précis
2. Une description détaillée expliquant pourquoi ce risque est pertinent
3. La catégorie et le type appropriés
4. Des estimations réalistes G, F, P basées sur le contexte

Ajoute aussi 2-3 recommandations générales pour le projet.

Répond STRICTEMENT en JSON avec ce format exact:
{{
  "project_title": "{request.analysis_title}",
  "identified_risks": [
    {{
      "title": "Titre du risque 1",
      "description": "Description détaillée du risque et pourquoi il s'applique à ce projet",
      "category": "Projet/Programme",
      "type": "Technique",
      "estimated_G": 4,
      "estimated_F": 3,
      "estimated_P": 2
    }},
    {{
      "title": "Titre du risque 2",
      "description": "Description détaillée...",
      "category": "Industriel",
      "type": "Cyber & SSI",
      "estimated_G": 5,
      "estimated_F": 2,
      "estimated_P": 3
    }}
  ],
  "general_recommendations": [
    "Recommandation générale 1 pour la gestion globale du projet",
    "Recommandation générale 2...",
    "Recommandation générale 3..."
  ]
}}
"""

    # Appel à l'IA
    llm_out = call_llm_for_risk(prompt, "Projet/Programme", "Technique", request.sector or "")

    if llm_out is None:
        raise HTTPException(status_code=503, detail="Service IA indisponible. Réessayez plus tard.")

    try:
        # Parser la réponse JSON
        import json
        response_text = llm_out.get("justification", "") or str(llm_out)
        # Nettoyer la réponse si nécessaire
        start = response_text.find("{")
        end = response_text.rfind("}") + 1
        if start != -1 and end != -1:
            json_str = response_text[start:end]
            parsed = json.loads(json_str)
        else:
            parsed = json.loads(response_text)

        # Valider la structure
        if not isinstance(parsed, dict) or "identified_risks" not in parsed:
            raise ValueError("Structure JSON invalide")

        # Convertir en objets Pydantic
        risks = []
        for risk_data in parsed["identified_risks"]:
            risk = IdentifiedRisk(**risk_data)
            risks.append(risk)

        recommendations = parsed.get("general_recommendations", [])

        return ProjectAnalysisResponse(
            project_title=request.analysis_title,
            identified_risks=risks,
            general_recommendations=recommendations
        )

    except Exception as e:
        print(f"Erreur parsing réponse IA: {e}")
        print(f"Réponse brute: {llm_out}")
        raise HTTPException(status_code=502, detail="Réponse IA malformée")
