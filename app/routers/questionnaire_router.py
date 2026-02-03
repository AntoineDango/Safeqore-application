from docx import Document
from app.services.llm_service import call_llm_for_risk
"""
Router pour le questionnaire (questions, analyse et historique)
Expose les endpoints attendus par le frontend mobile:
- GET /questionnaire/questions
- POST /questionnaire/analyze
- GET /questionnaire/analyses
- GET /questionnaire/analyses/{id}
"""

from fastapi import APIRouter, HTTPException, Query, Depends
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Literal
from datetime import datetime
from io import BytesIO
import os
import json
from app.auth.dependencies import get_current_user

# Stockage simple en JSON (même dossier que l'enrichissement)
BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
TRAINING_DATA_DIR = os.path.join(BASE_DIR, "training", "data")
ANALYSES_FILE = os.path.join(TRAINING_DATA_DIR, "questionnaire_analyses.json")

os.makedirs(TRAINING_DATA_DIR, exist_ok=True)

router = APIRouter(prefix="/questionnaire", tags=["Questionnaire"])

Dimension = Literal["G", "F", "P"]
Classification = Literal["Faible", "Modéré", "Élevé"]


class QuestionOption(BaseModel):
    id: str
    label: str
    contribution: Literal[1, 2, 3, 4, 5]
    niveau: Optional[Classification] = None


class Question(BaseModel):
    id: str
    dimension: Dimension
    texte_question: str
    reponses_possibles: List[QuestionOption]
    poids: int = 1
    secteurs_concernes: List[str] = []


class QuestionnaireQuestionsResponse(BaseModel):
    version: str
    questions: List[Question]


class AnswerItem(BaseModel):
    question_id: str
    option_id: str


class QuestionnaireAnalyzeRequest(BaseModel):
    description: str
    category: str
    type: str
    sector: Optional[str] = None
    answers: List[AnswerItem]


class QuestionnaireAnalyzeResponse(BaseModel):
    id: str
    timestamp: str
    questionnaire_version: str
    description: str
    category: str
    type: str
    sector: Optional[str]
    answers: List[AnswerItem]
    details: Dict[str, List[Dict[str, object]]]
    method: str
    G: int
    F: int
    P: int
    score: int
    classification: Classification
    justification: str
    normalized_score_100: int


class TraceListResponse(BaseModel):
    total: int
    limit: int
    offset: int
    items: List[QuestionnaireAnalyzeResponse]


# Residual (post-measure) models
class ResidualMeasure(BaseModel):
    text: str = Field(..., description="Description de la mesure")
    impacted: Dict[Literal["G","F","P"], bool] = Field(..., description="Facteurs impactés")
    new_values: Dict[Literal["G","F","P"], Optional[int]] = Field(default_factory=dict, description="Nouvelles valeurs proposées (1-5) pour les facteurs impactés")
    answers_by_dim: Dict[Literal["G","F","P"], List[AnswerItem]] = Field(default_factory=dict, description="Réponses (mini-questionnaire) par facteur impacté")


class ResidualRequest(BaseModel):
    parent_id: str = Field(..., description="ID de l'analyse d'origine")
    measures: List[ResidualMeasure]


class ResidualResponseItem(QuestionnaireAnalyzeResponse):
    parent_id: Optional[str] = None
    measure_text: Optional[str] = None


class ResidualResponse(BaseModel):
    items: List[ResidualResponseItem]


# Banque de questions (MVP). Versionnez-la pour que le frontend l'affiche.
QUESTIONNAIRE_VERSION = "1.0.0"

QUESTION_BANK: List[Question] = [
    # G - Gravité
    Question(
        id="G1",
        dimension="G",
        texte_question="Quel est l'impact potentiel en cas de matérialisation du risque ?",
        reponses_possibles=[
            QuestionOption(id="G1_O1", label="Impact mineur", contribution=1, niveau="Faible"),
            QuestionOption(id="G1_O2", label="Impact limité", contribution=2),
            QuestionOption(id="G1_O3", label="Impact modéré", contribution=3, niveau="Modéré"),
            QuestionOption(id="G1_O4", label="Impact majeur", contribution=4),
            QuestionOption(id="G1_O5", label="Impact critique", contribution=5, niveau="Élevé"),
        ],
        poids=1,
        secteurs_concernes=[],
    ),
    Question(
        id="G2",
        dimension="G",
        texte_question="Quelle est la criticité des actifs concernés (sécurité, finances, image) ?",
        reponses_possibles=[
            QuestionOption(id="G2_O1", label="Faible", contribution=1),
            QuestionOption(id="G2_O2", label="Légère", contribution=2),
            QuestionOption(id="G2_O3", label="Moyenne", contribution=3),
            QuestionOption(id="G2_O4", label="Élevée", contribution=4),
            QuestionOption(id="G2_O5", label="Très élevée", contribution=5),
        ],
        poids=1,
        secteurs_concernes=[],
    ),

    # F - Fréquence
    Question(
        id="F1",
        dimension="F",
        texte_question="À quelle fréquence l'exposition au risque a-t-elle lieu ?",
        reponses_possibles=[
            QuestionOption(id="F1_O1", label="Très rare", contribution=1, niveau="Faible"),
            QuestionOption(id="F1_O2", label="Occasionnelle", contribution=2),
            QuestionOption(id="F1_O3", label="Régulière", contribution=3, niveau="Modéré"),
            QuestionOption(id="F1_O4", label="Fréquente", contribution=4),
            QuestionOption(id="F1_O5", label="Quotidienne", contribution=5, niveau="Élevé"),
        ],
        poids=1,
        secteurs_concernes=[],
    ),
    Question(
        id="F2",
        dimension="F",
        texte_question="Quelle est la durée d'exposition au risque ?",
        reponses_possibles=[
            QuestionOption(id="F2_O1", label="Très courte", contribution=1),
            QuestionOption(id="F2_O2", label="Courte", contribution=2),
            QuestionOption(id="F2_O3", label="Modérée", contribution=3),
            QuestionOption(id="F2_O4", label="Longue", contribution=4),
            QuestionOption(id="F2_O5", label="Très longue", contribution=5),
        ],
        poids=1,
        secteurs_concernes=[],
    ),

    # P - Probabilité
    Question(
        id="P1",
        dimension="P",
        texte_question="Quelle est la probabilité d'occurrence de l'événement ?",
        reponses_possibles=[
            QuestionOption(id="P1_O1", label="Très improbable", contribution=1, niveau="Faible"),
            QuestionOption(id="P1_O2", label="Peu probable", contribution=2),
            QuestionOption(id="P1_O3", label="Assez probable", contribution=3, niveau="Modéré"),
            QuestionOption(id="P1_O4", label="Probable", contribution=4),
            QuestionOption(id="P1_O5", label="Très probable", contribution=5, niveau="Élevé"),
        ],
        poids=1,
        secteurs_concernes=[],
    ),
    Question(
        id="P2",
        dimension="P",
        texte_question="Existe-t-il des vulnérabilités connues augmentant la probabilité ?",
        reponses_possibles=[
            QuestionOption(id="P2_O1", label="Aucune vulnérabilité", contribution=1),
            QuestionOption(id="P2_O2", label="Peu de vulnérabilités", contribution=2),
            QuestionOption(id="P2_O3", label="Quelques vulnérabilités", contribution=3),
            QuestionOption(id="P2_O4", label="Plusieurs vulnérabilités", contribution=4),
            QuestionOption(id="P2_O5", label="Nombreuses vulnérabilités", contribution=5),
        ],
        poids=1,
        secteurs_concernes=[],
    ),
]


# Helpers de stockage

def _load_analyses() -> List[dict]:
    if os.path.exists(ANALYSES_FILE):
        try:
            with open(ANALYSES_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return []
    return []


def _save_analyses(data: List[dict]):
    with open(ANALYSES_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)


# Calcul Kinney simple

def _classify(score: int) -> Classification:
    if score <= 25:
        return "Faible"
    elif score <= 50:
        return "Modéré"
    return "Élevé"


@router.get("/questions", response_model=QuestionnaireQuestionsResponse)
async def get_questions(sector: Optional[str] = Query(None, description="Filtrer par secteur (optionnel)")):
    if sector:
        filtered = [q for q in QUESTION_BANK if not q.secteurs_concernes or sector in q.secteurs_concernes]
    else:
        filtered = QUESTION_BANK
    return QuestionnaireQuestionsResponse(version=QUESTIONNAIRE_VERSION, questions=filtered)


@router.post("/analyze", response_model=QuestionnaireAnalyzeResponse)
async def analyze_questionnaire(
    payload: QuestionnaireAnalyzeRequest,
    current_user: dict = Depends(get_current_user)
):
    user_uid = current_user["uid"]
    user_email = current_user.get("email", "")
    print(f"[QuestionnaireRouter] Creating analysis for user: {user_uid} ({user_email})")
    
    # Indexation rapide des questions et options
    q_map: Dict[str, Question] = {q.id: q for q in QUESTION_BANK}
    opt_map: Dict[str, Dict[str, QuestionOption]] = {
        q.id: {o.id: o for o in q.reponses_possibles} for q in QUESTION_BANK
    }

    # Agréger contributions par dimension (pondérées)
    sums = {"G": 0.0, "F": 0.0, "P": 0.0}
    weights = {"G": 0.0, "F": 0.0, "P": 0.0}
    details: Dict[str, List[Dict[str, object]]] = {"G": [], "F": [], "P": []}

    for ans in payload.answers:
        q = q_map.get(ans.question_id)
        if not q:
            # ignorer silencieusement ou lever une erreur? Ici on ignore.
            continue
        o = opt_map[q.id].get(ans.option_id)
        if not o:
            continue
        w = float(q.poids or 1)
        sums[q.dimension] += o.contribution * w
        weights[q.dimension] += w
        details[q.dimension].append({
            "question_id": q.id,
            "option_id": o.id,
            "contribution": o.contribution,
            "poids": q.poids,
        })

    def _compute(dim: str) -> int:
        if weights[dim] <= 0:
            return 1  # défaut minimal si aucune réponse
        val = round(sums[dim] / weights[dim])
        return max(1, min(5, int(val)))

    G = _compute("G")
    F = _compute("F")
    P = _compute("P")
    score = int(G * F * P)
    normalized_score_100 = int(round(score / 125 * 100))
    classification = _classify(score)

    # Construire la réponse
    all_data = _load_analyses()
    new_id = f"QA_{datetime.now().strftime('%Y%m%d%H%M%S')}_{len(all_data)}"
    item = {
        "id": new_id,
        "timestamp": datetime.now().isoformat(),
        "questionnaire_version": QUESTIONNAIRE_VERSION,
        "description": payload.description,
        "category": payload.category,
        "type": payload.type,
        "sector": payload.sector or "",
        "answers": [a.model_dump() for a in payload.answers],
        "details": details,
        "method": "questionnaire",
        "G": G,
        "F": F,
        "P": P,
        "score": score,
        "classification": classification,
        "justification": "Calcul basé sur la moyenne pondérée des contributions par dimension (méthode Kinney).",
        "normalized_score_100": normalized_score_100,
        "user_uid": user_uid,  # IMPORTANT: Lier l'analyse à l'utilisateur
        "user_email": user_email,
    }
    all_data.append(item)
    _save_analyses(all_data)
    
    print(f"[QuestionnaireRouter] Analysis {new_id} created successfully for user {user_uid}")

    return QuestionnaireAnalyzeResponse(**item)  # type: ignore[arg-type]


@router.get("/analyses", response_model=TraceListResponse)
async def list_analyses(limit: int = 50, offset: int = 0):
    data = _load_analyses()
    total = len(data)
    items = data[offset: offset + limit]
    return TraceListResponse(total=total, limit=limit, offset=offset, items=items)


@router.get("/analyses/{analysis_id}")
async def get_analysis(analysis_id: str):
    data = _load_analyses()
    for it in data:
        if it.get("id") == analysis_id:
            return it
    raise HTTPException(status_code=404, detail=f"Analyse {analysis_id} non trouvée")


@router.post("/residual", response_model=ResidualResponse)
async def create_residual(payload: ResidualRequest):
    data = _load_analyses()
    parent = next((x for x in data if x.get("id") == payload.parent_id), None)
    if not parent:
        raise HTTPException(status_code=404, detail="Analyse d'origine introuvable")

    orig_G = int(parent.get("G", 1))
    orig_F = int(parent.get("F", 1))
    orig_P = int(parent.get("P", 1))

    items: List[ResidualResponseItem] = []
    for idx, m in enumerate(payload.measures):
        imp = m.impacted or {"G": False, "F": False, "P": False}
        # validation: au moins un facteur impacté
        if not (imp.get("G") or imp.get("F") or imp.get("P")):
            raise HTTPException(status_code=400, detail=f"Mesure #{idx+1}: aucun facteur impacté")

        # Calcul via mini-questionnaire pour chaque dimension impactée
        def _compute_from_answers(dim: str, answers: List[AnswerItem]) -> int:
            rel_q = {q.id: q for q in QUESTION_BANK if q.dimension == dim}
            sums = 0.0
            weights = 0.0
            for a in answers:
                q = rel_q.get(a.question_id)
                if not q:
                    continue
                opt = next((o for o in q.reponses_possibles if o.id == a.option_id), None)
                if not opt:
                    continue
                w = float(q.poids or 1)
                sums += opt.contribution * w
                weights += w
            if weights <= 0:
                raise HTTPException(status_code=400, detail=f"Mesure #{idx+1}: réponses invalides pour {dim}")
            val = round(sums / weights)
            return max(1, min(5, int(val)))

        ab = getattr(m, "answers_by_dim", {}) or {}

        if imp.get("G", False):
            g_answers = ab.get("G", [])  # type: ignore[index]
            if not g_answers:
                raise HTTPException(status_code=400, detail=f"Mesure #{idx+1}: réponses questionnaire requises pour G")
            G = _compute_from_answers("G", g_answers)
        else:
            G = orig_G

        if imp.get("F", False):
            f_answers = ab.get("F", [])  # type: ignore[index]
            if not f_answers:
                raise HTTPException(status_code=400, detail=f"Mesure #{idx+1}: réponses questionnaire requises pour F")
            F = _compute_from_answers("F", f_answers)
        else:
            F = orig_F

        if imp.get("P", False):
            p_answers = ab.get("P", [])  # type: ignore[index]
            if not p_answers:
                raise HTTPException(status_code=400, detail=f"Mesure #{idx+1}: réponses questionnaire requises pour P")
            P = _compute_from_answers("P", p_answers)
        else:
            P = orig_P

        score = int(G * F * P)
        classification = _classify(score)
        normalized_score_100 = int(round(score / 125 * 100))

        new_id = f"QR_{datetime.now().strftime('%Y%m%d%H%M%S')}_{len(data)}"
        item = {
            "id": new_id,
            "timestamp": datetime.now().isoformat(),
            "questionnaire_version": QUESTIONNAIRE_VERSION,
            "description": parent.get("description"),
            "category": parent.get("category"),
            "type": parent.get("type"),
            "sector": parent.get("sector"),
            "answers": parent.get("answers", []),
            "details": {"G": [], "F": [], "P": []},
            "method": "residual",
            "G": G,
            "F": F,
            "P": P,
            "score": score,
            "classification": classification,
            "normalized_score_100": normalized_score_100,
            "justification": "Ré-estimation post-mesure (facteurs non impactés conservés)",
            "parent_id": payload.parent_id,
            "measure_text": m.text,
        }
        data.append(item)
        items.append(ResidualResponseItem(**item))

    _save_analyses(data)
    return ResidualResponse(items=items)


@router.get("/report/{analysis_id}")
async def generate_report(analysis_id: str):
    data = _load_analyses()
    main = next((x for x in data if x.get("id") == analysis_id), None)
    if not main:
        raise HTTPException(status_code=404, detail="Analyse introuvable")

    # Collect residuals linked to this analysis
    residuals = [x for x in data if x.get("method") == "residual" and x.get("parent_id") == analysis_id]

    # IA analysis (best-effort)
    ia = None
    try:
        ia_out = call_llm_for_risk(main.get("description"), main.get("category"), main.get("type"), main.get("sector"))
        if ia_out is not None:
            ia_G = int(ia_out.get("G"))
            ia_F = int(ia_out.get("F"))
            ia_P = int(ia_out.get("P"))
            ia_score = int(ia_G * ia_F * ia_P)
            ia_cls = _classify(ia_score)
            ia = {"G": ia_G, "F": ia_F, "P": ia_P, "score": ia_score, "classification": ia_cls}
    except Exception:
        ia = None

    # Build Word document
    doc = Document()
    doc.add_heading("Rapport d'analyse de risque", level=0)
    doc.add_paragraph(f"ID: {main.get('id')}")
    doc.add_paragraph(f"Date: {main.get('timestamp')}")

    doc.add_heading("Contexte", level=1)
    doc.add_paragraph(f"Description: {main.get('description')}")
    doc.add_paragraph(f"Catégorie: {main.get('category')}  |  Type: {main.get('type')}  |  Secteur: {main.get('sector')}")

    doc.add_heading("Analyse utilisateur", level=1)
    doc.add_paragraph(f"Méthode: {main.get('method')}")
    doc.add_paragraph(f"G: {main.get('G')}  F: {main.get('F')}  P: {main.get('P')}")
    doc.add_paragraph(f"Score: {main.get('normalized_score_100')} / 100")
    doc.add_paragraph(f"Classification: {main.get('classification')}")
    
    # Graphique visuel du score
    def score_bar(val: int, max_val: int = 100, length: int = 50) -> str:
        pct = max(0, min(max_val, val)) / float(max_val)
        filled = int(round(pct * length))
        return "█" * filled + "░" * (length - filled)
    
    score_normalized = main.get('normalized_score_100', 0)
    doc.add_paragraph(f"Visualisation: {score_bar(score_normalized)}")
    
    # Graphique des facteurs G, F, P
    doc.add_heading("Détail des facteurs", level=2)
    
    def factor_bar(val: int, max_val: int = 5, length: int = 20) -> str:
        pct = max(0, min(max_val, val)) / float(max_val)
        filled = int(round(pct * length))
        return "█" * filled + "░" * (length - filled)
    
    factors_table = doc.add_table(rows=4, cols=2)
    factors_table.style = "Table Grid"
    factors_table.cell(0,0).text = "Facteur"
    factors_table.cell(0,1).text = "Valeur (sur 5)"
    
    factors_table.cell(1,0).text = "G (Gravité)"
    factors_table.cell(1,1).text = f"{main.get('G')}/5  |  {factor_bar(int(main.get('G', 1)))}"
    
    factors_table.cell(2,0).text = "F (Fréquence)"
    factors_table.cell(2,1).text = f"{main.get('F')}/5  |  {factor_bar(int(main.get('F', 1)))}"
    
    factors_table.cell(3,0).text = "P (Probabilité)"
    factors_table.cell(3,1).text = f"{main.get('P')}/5  |  {factor_bar(int(main.get('P', 1)))}"

    if ia:
        doc.add_heading("Analyse IA (assistance)", level=1)
        ia_score_normalized = int(round(ia['score'] / 125 * 100))
        doc.add_paragraph(f"G: {ia['G']}  F: {ia['F']}  P: {ia['P']}")
        doc.add_paragraph(f"Score: {ia_score_normalized} / 100  |  Classification: {ia['classification']}")
        doc.add_paragraph(f"Visualisation: {score_bar(ia_score_normalized)}")

        # Comparison
        doc.add_heading("Comparaison (Utilisateur vs IA)", level=1)
        try:
            h_score_norm = int(main.get("normalized_score_100", 0))
            diff = abs(h_score_norm - ia_score_normalized)
            doc.add_paragraph(f"Écart de score: {diff} points (sur 100)")
            
            # Tableau comparatif
            comp_table = doc.add_table(rows=3, cols=3)
            comp_table.style = "Table Grid"
            comp_table.cell(0,0).text = ""
            comp_table.cell(0,1).text = "Utilisateur"
            comp_table.cell(0,2).text = "IA"
            
            comp_table.cell(1,0).text = "Score /100"
            comp_table.cell(1,1).text = str(h_score_norm)
            comp_table.cell(1,2).text = str(ia_score_normalized)
            
            comp_table.cell(2,0).text = "Classification"
            comp_table.cell(2,1).text = str(main.get('classification', ''))
            comp_table.cell(2,2).text = str(ia['classification'])
            
            doc.add_paragraph("Les décisions finales appartiennent à l'utilisateur.")
        except Exception:
            pass

    # Residuals
    doc.add_heading("Mesures et risques résiduels", level=1)
    if residuals:
        for r in residuals:
            p = doc.add_paragraph()
            p.add_run("Mesure: ").bold = True
            p.add_run(str(r.get("measure_text")))
            doc.add_paragraph(f"Après mesure → G:{r.get('G')} F:{r.get('F')} P:{r.get('P')} Score:{r.get('score')} ({r.get('classification')})")
    else:
        doc.add_paragraph("Aucune mesure enregistrée")

    # Simple risk matrix table (G x P)
    doc.add_heading("Matrice de risques (G × P)", level=1)
    try:
        G = int(main.get("G", 1))
        P = int(main.get("P", 1))
        table = doc.add_table(rows=6, cols=6)
        table.style = "Table Grid"
        # headers
        table.cell(0,0).text = "G\\P"
        for j in range(1,6):
            table.cell(0,j).text = str(j)
        for i in range(1,6):
            table.cell(i,0).text = str(6 - i)  # G from 5 to 1
        # Mark point
        g_row = 6 - G
        p_col = P
        table.cell(g_row, p_col).text = "X"
    except Exception:
        pass

    # Stream as response
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    filename = f"rapport_{analysis_id}.docx"
    headers = {"Content-Disposition": f"attachment; filename={filename}"}
    return StreamingResponse(bio, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers=headers)


@router.get("/export")
async def export_analyses():
    """Export all analyses as JSON attachment."""
    data = _load_analyses()
    bio = BytesIO()
    bio.write(json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8"))
    bio.seek(0)
    headers = {"Content-Disposition": "attachment; filename=analyses_export.json"}
    return StreamingResponse(bio, media_type="application/json", headers=headers)


class ImportPayload(BaseModel):
    items: List[dict]


@router.post("/import")
async def import_analyses(payload: ImportPayload):
    """Import analyses: merge unique by id (skip duplicates)."""
    if not isinstance(payload.items, list):
        raise HTTPException(status_code=400, detail="items doit être une liste")
    data = _load_analyses()
    existing_ids = {x.get("id") for x in data}
    imported = 0
    for it in payload.items:
        iid = it.get("id")
        if not iid or iid in existing_ids:
            continue
        data.append(it)
        existing_ids.add(iid)
        imported += 1
    _save_analyses(data)
    return {"status": "ok", "imported": imported, "total": len(data)}
